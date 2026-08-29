"""書類生成(Office様式の差し込み) — Office形式の成果物を組み立てる。

重説・マイソクは実務でExcel/Word固定様式として作る。HTMLで近似を手書きしない。
- マイソク: blank workbookから自作した原紙(`templates/office/maisoku_genshi.xlsx`)に差し込み→.xlsx
- 重説: 35条書面のWord様式を生成・差し込み→.docx (build_juusetsu_docx)
PDF化は LibreOffice(`soffice --headless --convert-to pdf`)で別途(office_to_pdf)。
依存: openpyxl / python-docx（Office出力時だけ。requirements-office.txtで固定）。
"""
from __future__ import annotations

import shutil
import subprocess
import tempfile
import re
import sys
import os
import base64
import binascii
import io
import zipfile
from copy import copy
from pathlib import Path

_TPL = Path(__file__).parent / "templates" / "office"
MAISOKU_TEMPLATES = {
    "A": _TPL / "maisoku_a_standard_landscape.xlsx",
    "B": _TPL / "maisoku_b_buyer_portrait.xlsx",
    "C": _TPL / "maisoku_c_fax_mono.xlsx",
}
# Backward-compatible name used by existing callers/tests; A is the canonical default.
MAISOKU_GENSHI = MAISOKU_TEMPLATES["A"]
_ILLEGAL_XLSX_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")

# 原紙(マイソク原紙シート)のT列ラベル → maisoku.MAISOKU_FIELDS のキー
_GENSHI_MAP = {
    "物件所在地": "address", "交通": "access", "敷地面積": "land_area",
    "土地権利": "land_right", "地目": "chimoku", "建物構造": "structure",
    "延床面積": "building_area", "築年月": "built", "都市計画": "toshi_keikaku",
    "用途地域": "youto", "防火規制": "bouka", "他の法令上の制限": "other_law",
    "接道": "road", "設備": "setsubi", "総戸数": "total_units",
    "間取り": "floor_plan", "現況": "genkyo", "引渡し": "hikiwatashi",
}


def _norm(s) -> str:
    return str(s or "").replace("　", "").replace(" ", "").strip()


def fill_maisoku_xlsx(fields: dict, out_path, *, template: Path | None = None,
                      company: dict | None = None) -> str:
    return fill_maisoku_xlsx_with_profile(fields, out_path, template=template, company=company)


_XL_FONT_NAMES = {
    "gothic": "Yu Gothic",
    "condensed": "Yu Gothic UI",
    "rounded": "Meiryo",
}


def _defined_destinations(wb, prefix: str, *, exact: bool = False):
    """Yield (worksheet, coordinate) for matching defined names."""
    for name, defined in wb.defined_names.items():
        matches = name == prefix if exact else name.startswith(prefix)
        if not matches:
            continue
        for sheet_name, coordinate in defined.destinations:
            if sheet_name in wb.sheetnames:
                yield wb[sheet_name], coordinate.replace("$", "")


def _named_cells(wb, prefix: str, *, exact: bool = False):
    for ws, coordinate in _defined_destinations(wb, prefix, exact=exact):
        selected = ws[coordinate]
        if isinstance(selected, tuple):
            for row in selected:
                if isinstance(row, tuple):
                    yield from row
                else:
                    yield row
        else:
            yield selected


def _write_named_fields(wb, fields: dict) -> bool:
    """Write the clean-room template contract (ms_<field> defined names)."""
    names = set(wb.defined_names)
    mapped = {name for name in names if name.startswith("ms_")}
    if not mapped:
        return False
    from hub_core import maisoku as _maisoku

    expected = {"ms_" + key for key in _maisoku.FIELD_KEYS}
    missing = sorted(expected - mapped)
    if missing:
        raise ValueError(
            "Excel原紙のnamed rangeが不足しています: "
            + "、".join(name[3:] for name in missing)
        )
    f = dict(fields or {})
    f["price_label"] = str(f.get("price_label") or "価格")
    for name in mapped:
        key = name[3:]
        value = f.get(key, "")
        targets = list(_named_cells(wb, name, exact=True))
        if len(targets) != 1:
            raise ValueError(f"Excel原紙の項目 {key} の書込先が一意ではありません。")
        _set_excel_text(targets[0], value)
    return True


def _set_excel_text(cell, value) -> None:
    """Keep untrusted document fields as text, including strings starting with '='."""
    cell.value = _clean_excel_text(value)
    cell.data_type = "s"


def _clean_excel_text(value) -> str:
    return _ILLEGAL_XLSX_CHARS.sub("", "" if value is None else str(value))


def _apply_named_theme(wb, fields: dict, company: dict | None) -> None:
    """Apply the document's pinned tenant color/font to named style ranges."""
    from hub_core import maisoku as _maisoku

    profile = dict(company or {})
    accent = str((fields or {}).get("_accent") or profile.get("brand_color") or "#2A2E37")
    font_key = str((fields or {}).get("_font") or profile.get("display_font") or "gothic")
    tokens = _maisoku.theme_fields(profile, accent=accent, font=font_key)
    colors = {
        "style_accent_fill_": tokens["accent"],
        "style_accent_ink_": tokens["accent_ink"],
        "style_accent_text_": tokens["accent_text"],
        "style_accent_soft_": tokens["accent_soft"],
    }

    # Body first, then title ranges override it.
    for prefix in ("style_font_body_", "style_font_title_"):
        for cell in _named_cells(wb, prefix):
            font = copy(cell.font)
            font.name = _XL_FONT_NAMES.get(font_key, _XL_FONT_NAMES["gothic"])
            cell.font = font
    for prefix, color in colors.items():
        rgb = "FF" + str(color).lstrip("#").upper()
        for cell in _named_cells(wb, prefix):
            if prefix in ("style_accent_fill_", "style_accent_soft_"):
                from openpyxl.styles import PatternFill
                cell.fill = PatternFill("solid", fgColor=rgb)
            else:
                font = copy(cell.font)
                font.color = rgb
                cell.font = font


def _xlsx_variant(fields: dict) -> str:
    raw = str((fields or {}).get("_xlsx_variant") or "").strip().upper()
    if raw in MAISOKU_TEMPLATES:
        return raw
    legacy = str((fields or {}).get("_variant") or "").strip()
    if legacy == "v2-tate":
        return "B"
    if legacy == "v2-sumi":
        return "C"
    return "A"


def _apply_xlsx_variant_style(wb, variant: str) -> None:
    """C/FAXは会社色の有無にかかわらず再現可能な白黒へ固定する。"""
    if variant != "C":
        return
    from openpyxl.styles import Border, PatternFill, Side

    black = "FF111111"
    white = "FFFFFFFF"
    gray = "FFE6E6E6"
    rule = Side(style="thin", color="FF777777")
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                dark_fill = False
                if cell.fill.fill_type == "solid":
                    rgb = str(getattr(cell.fill.fgColor, "rgb", "") or "")[-6:]
                    try:
                        red, green, blue = (int(rgb[i:i + 2], 16) for i in (0, 2, 4))
                        dark_fill = (red * 299 + green * 587 + blue * 114) < 150000
                    except ValueError:
                        dark_fill = False
                    cell.fill = PatternFill("solid", fgColor=("FF222222" if dark_fill else gray))
                font = copy(cell.font)
                font.color = white if dark_fill else black
                cell.font = font
                if any(side.style for side in (cell.border.left, cell.border.right,
                                                cell.border.top, cell.border.bottom)):
                    cell.border = Border(left=rule, right=rule, top=rule, bottom=rule)


def _decoded_image(value: str, *, grayscale: bool):
    """Return a verified in-memory PNG/JPEG payload; never fetch paths or URLs."""
    raw = str(value or "").strip()
    match = re.fullmatch(r"data:image/(png|jpeg|webp);base64,([A-Za-z0-9+/=\r\n]+)", raw, re.I)
    if not match:
        return None
    try:
        payload = base64.b64decode(match.group(2), validate=True)
    except (ValueError, binascii.Error):
        return None
    if not payload or len(payload) > 20 * 1024 * 1024:
        return None
    try:
        from PIL import Image as PILImage, ImageOps
        source = PILImage.open(io.BytesIO(payload))
        source.load()
        if source.width < 1 or source.height < 1 or source.width * source.height > 80_000_000:
            return None
        if grayscale:
            source = ImageOps.grayscale(source)
        out = io.BytesIO()
        if source.mode not in ("RGB", "RGBA"):
            source = source.convert("RGB")
        source.save(out, format="PNG")
        out.seek(0)
        return out
    except (ImportError, OSError, ValueError):
        return None


def _embed_maisoku_images(wb, fields: dict, variant: str) -> list[io.BytesIO]:
    """Embed main photo/floor plan into OOXML media parts; no external relationship."""
    from openpyxl.drawing.image import Image as XLImage

    streams: list[io.BytesIO] = []
    max_width = 350 if variant == "B" else 430
    max_height = 155
    for field, anchor_name in (("photo_main", "image_photo_main"),
                               ("photo_floorplan", "image_floorplan")):
        stream = _decoded_image((fields or {}).get(field, ""), grayscale=(variant == "C"))
        if stream is None:
            continue
        destinations = list(_defined_destinations(wb, anchor_name, exact=True))
        if len(destinations) != 1:
            continue
        ws, coordinate = destinations[0]
        image = XLImage(stream)
        scale = min(max_width / image.width, max_height / image.height, 1.0)
        image.width = max(1, round(image.width * scale))
        image.height = max(1, round(image.height * scale))
        ws.add_image(image, coordinate.split(":", 1)[0])
        streams.append(stream)
    return streams


def _fill_legacy_maisoku(ws, fields: dict) -> None:
    """Compatibility for user-supplied templates using the pre-2026 label layout."""
    f = fields or {}
    for r in range(1, ws.max_row + 1):
        label = _norm(ws.cell(row=r, column=20).value)
        if not label:
            continue
        key = _GENSHI_MAP.get(label)
        if key and str(f.get(key, "")).strip():
            _set_excel_text(ws.cell(row=r, column=23), f.get(key))
        elif label in ("建ぺい率", "容積率"):
            parts = str(f.get("kenpei_yoseki", "") or "").replace("／", "/").split("/")
            if label == "建ぺい率" and parts and parts[0].strip():
                _set_excel_text(ws.cell(row=r, column=23), parts[0].strip())
            elif label == "容積率" and len(parts) > 1:
                _set_excel_text(ws.cell(row=r, column=23), parts[1].strip())
    price = str(f.get("price", "") or "").replace("万円", "").strip()
    if price:
        _set_excel_text(ws["U8"], price)
    name = (f.get("property_name") or f.get("title_copy") or "").strip()
    if name:
        yld = (f.get("yield_rate") or "").strip()
        _set_excel_text(ws["B3"], f"{name}　利回り{yld}" if yld else name)
    for coord, val in {
        "E40": f.get("company_name"), "E39": f.get("license"),
        "M39": f.get("company_address"), "M40": _tel_fax(f),
        "M41": ("メール：" + f["company_email"]) if f.get("company_email") else "",
        "T41": ("　　　担当：" + f["staff"]) if f.get("staff") else "",
    }.items():
        if val and str(val).strip():
            _set_excel_text(ws[coord], val)
    try:
        ws._images = [im for im in getattr(ws, "_images", [])
                      if getattr(getattr(im, "anchor", None), "_from", None) is None
                      or im.anchor._from.col >= 5]
    except Exception:
        pass


def fill_maisoku_xlsx_with_profile(fields: dict, out_path, *, template: Path | None = None,
                                    company: dict | None = None) -> str:
    """Fill the original Excel template and apply the pinned tenant profile."""
    variant = _xlsx_variant(fields)
    _tpl = template or MAISOKU_TEMPLATES[variant]
    if not Path(_tpl).is_file():
        raise FileNotFoundError("Excel原紙テンプレートがありません。配布物が欠けています。")
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("Excel出力には openpyxl が必要です。DEMO.md の手順で追加してください。") from exc
    src = Path(_tpl)
    wb = openpyxl.load_workbook(src, keep_links=False)
    ws = wb["マイソク原紙"] if "マイソク原紙" in wb.sheetnames else wb.worksheets[0]
    # User-supplied templates can carry sample sheets. Output only the selected sheet.
    for name in list(wb.sheetnames):
        if name != ws.title:
            del wb[name]
    from hub_core import maisoku as _maisoku

    f = _maisoku.fields_with_company(dict(fields or {}), company or {})
    if _write_named_fields(wb, f):
        _apply_named_theme(wb, f, company)
        _apply_xlsx_variant_style(wb, variant)
        image_streams = _embed_maisoku_images(wb, f, variant)
        missing = _maisoku.check_required(f)
        ws["A37"] = ("掲載情報 / 下書き（必要表示事項に未入力あり）"
                     if missing else "掲載情報")
    else:
        image_streams = []
        _fill_legacy_maisoku(ws, f)
    wb._external_links = []
    tenant_name = _clean_excel_text(
        f.get("company_name") or (company or {}).get("name") or "").strip()
    if tenant_name:
        wb.properties.creator = tenant_name
        wb.properties.lastModifiedBy = tenant_name
    wb.properties.title = _clean_excel_text(f.get("property_name") or "マイソク")

    out = str(Path(out_path).expanduser())
    wb.save(out)
    for stream in image_streams:
        stream.close()
    return out


def _tel_fax(f) -> str:
    tel = (f.get("company_tel") or "").strip()
    fax = (f.get("company_fax") or "").strip()
    parts = []
    if tel:
        parts.append("TEL : " + tel)
    if fax:
        parts.append("FAX : " + fax)
    return "　".join(parts)


# 重説(35条)の完全な項目体系の正本。Workflow(国交省標準様式+建物賃貸様式の構造抽出＋35条法定事項
# 調査＋敵対的完全性ループ)で確定。売買=baibai(29章213項目)/賃貸=chintai(27章129項目)。
JUUSETSU_SCHEMA = _TPL / "juusetsu_schema.json"

import re as _re

# 項目ラベル末尾の冗長な「根拠・確認導線」括弧(完全性検証エージェントの註記)を除去し、フォーム向けに整える
_META_PAREN = _re.compile(
    r"（[^（）]*(?:確認item|確認導線|独立確認|根拠|実務|漏れ|非対称|レンズ|防ぐ|前提|連動|母数|"
    r"直結|留まり|解消|併設|item|導線|事故|取りこぼ|誤認|連携|拘束|前段)[^（）]*）")


def _clean_item(s: str) -> str:
    s = str(s or "").strip()
    s = _META_PAREN.sub("", s)
    if len(s) > 80:
        s = _re.split(r"[、，]", s, maxsplit=1)[0]
    return s.strip().rstrip("、。 ")[:100]


def juusetsu_sections(deal_type: str = "売買", property_kind: str = ""):
    """deal_type に応じた(sections, key)。様式選択は deal_taxonomy の正規形 transaction 経由に一本化。

    property_kind が確定していれば、区分所有追加章など applies_property_kinds に該当しない章を落とす
    （未確定=None のときは落とさない＝現行どおり全章）。render_juusetsu_md(A) と同一判定。
    """
    import json as _j

    from . import deal_taxonomy as _tax
    schema = _j.loads(Path(JUUSETSU_SCHEMA).read_text(encoding="utf-8"))
    key = "chintai" if _tax.normalize_transaction(deal_type) == "lease" else "baibai"
    pk = _tax.normalize_property_kind(property_kind)
    secs = [s for s in schema.get(key, []) if _tax.section_applies(s.get("applies_property_kinds"), pk)]
    return secs, key


def build_md_docx(md_text: str, out_path, *, prs_note: str = "",
                  document_status: str = "", audit_label: str = "") -> str:
    """markdown（見出し/表/箇条書き）→ .docx。重説/精算書の下書き本文を客に渡せるWordにする。
    プレビューと同一内容を出力する（build_juusetsu_docx({})の空様式バグ是正）。"""
    import re as _re
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    def _plain(s: str) -> str:
        s = _re.sub(r"\*\*(.+?)\*\*", r"\1", s)   # 太字マーカ除去
        return s.strip()

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)
    if document_status in ("draft", "final"):
        is_final = document_status == "final"
        status_text = ("確定版（監査照合済み）" if is_final
                       else "DRAFT / 下書き・交付不可")
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run(status_text + ((" / " + audit_label) if audit_label else ""))
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x16, 0x45, 0x2A) if is_final else RGBColor(0xB4, 0x23, 0x18)
        for section in doc.sections:
            header = section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header.add_run(status_text)
            header_run.bold = True
            header_run.font.color.rgb = run.font.color.rgb
    lines = (md_text or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.strip().startswith("<!-- ainote-juusetsu-schema:") and line.strip().endswith("-->"):
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(_plain(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(_plain(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(_plain(line[4:]), level=2)
        elif line.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [_plain(c) for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):   # 区切り行(|---|)は除外
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                try:
                    tbl.style = "Table Grid"
                except KeyError:
                    pass
                for ri, r in enumerate(rows):
                    for ci in range(ncol):
                        tbl.cell(ri, ci).text = r[ci] if ci < len(r) else ""
            continue
        elif line.startswith(">"):
            para = doc.add_paragraph(_plain(line.lstrip(">").strip()))
            para.runs[0].italic = True if para.runs else None
        elif line.lstrip().startswith("- "):
            doc.add_paragraph(_plain(line.lstrip()[2:]), style="List Bullet")
        else:
            doc.add_paragraph(_plain(line))
        i += 1
    if prs_note:
        doc.add_paragraph().add_run("◆ 災害リスク事前スクリーニング（参考）：" + prs_note)
    doc.save(str(out_path))
    return str(out_path)


def build_juusetsu_docx(data: dict, out_path, *, deal_type: str = "売買", property_kind: str = "", prs_note: str = "") -> str:
    """重要事項説明書(35条)を実物様式どおりの完全な章立てで .docx 生成。
    売買と賃貸で項目体系が異なる(deal_type)。property_kind で区分所有追加章を出し分ける。
    各項目は宅建士が登記簿/公的資料/現地で確認し記入。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    d = data or {}
    sections, key = juusetsu_sections(deal_type, property_kind)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Mincho"
    style.font.size = Pt(8.5)
    try:
        from docx.oxml.ns import qn
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Mincho")
    except Exception:
        pass
    for sec in doc.sections:  # 多項目をA4に収めるため余白を狭く
        sec.top_margin = sec.bottom_margin = Pt(34)
        sec.left_margin = sec.right_margin = Pt(38)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("重要事項説明書（" + ("建物賃貸借用" if key == "chintai" else "売買・交換") + "）")
    r.bold = True
    r.font.size = Pt(15)
    doc.add_paragraph().add_run(
        "下記の不動産について、宅地建物取引業法第35条の規定に基づき、次のとおり説明します。"
        "この内容は重要ですから、十分理解されるようお願いします。").font.size = Pt(8.5)

    # PRS災害リスク事前スクリーニング(参考)。法定ハザードマップ説明の代替でない旨はprs_note本文に明記済。
    if prs_note:
        box = doc.add_paragraph()
        box.paragraph_format.space_before = Pt(4)
        box.paragraph_format.space_after = Pt(4)
        br = box.add_run("◆ 災害リスク事前スクリーニング（参考）：" + prs_note)
        br.font.size = Pt(8)
        br.italic = True

    for sc in sections:
        head = (str(sc.get("no", "")).strip() + "　" + str(sc.get("title", "")).strip()).strip()
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        hr = h.add_run(head)
        hr.bold = True
        hr.font.size = Pt(10)
        items = sc.get("items", [])
        if not items:
            continue
        tbl = doc.add_table(rows=len(items), cols=2)
        tbl.style = "Table Grid"
        for i, it in enumerate(items):
            tbl.rows[i].cells[0].text = _clean_item(it)
            val = d.get(_clean_item(it)) or d.get(it) or ""
            tbl.rows[i].cells[1].text = str(val)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.add_run("※ 本書は下書きの雛形です。各欄は登記簿・公的資料・現地で確認のうえ記入し、"
                 "宅地建物取引士が宅地建物取引士証を提示して説明・記名（押印）し交付してください。"
                 "説明の最終責任は記名した宅地建物取引士に帰属します。").font.size = Pt(8)

    out = str(Path(out_path).expanduser())
    doc.save(out)
    return out


# 37条書面（契約成立時交付書面）の法定記載事項（宅建業法37条・売買/賃貸で一部差異）。
# 数値・金額・法令の確定はしない＝雛形の枠と確認欄を出し、記入と記名は人間（宅建士）。
KEIYAKU37_ITEMS_COMMON = [
    "当事者の氏名・住所",
    "宅地建物を特定するために必要な表示（所在・地番・地目・面積・建物の構造/種類等）",
    "代金・交換差金・借賃の額、支払時期・方法",
    "宅地建物の引渡しの時期",
    "移転登記の申請時期",  # 売買のみ（賃貸は該当なし）
    "契約の解除に関する定め",
    "損害賠償額の予定・違約金に関する定め",
    "代金・交換差金・借賃以外の金銭の授受に関する定め（額・目的・授受時期）",
    "天災その他不可抗力による損害の負担（危険負担）に関する定め",
    "租税その他の公課の負担に関する定め",
    "契約不適合責任（担保責任）／その履行に関する保証保険契約等の定め",
    "ローン・代金等の金銭の貸借のあっせんに関する定め、あっせん不成立時の措置",
]


_DETERMINISTIC_DOCX_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_DETERMINISTIC_DOCX_MODE = 0o600 << 16


def _normalize_docx_package(path: str | Path) -> None:
    """Rewrite a python-docx OPC package with deterministic ZIP metadata/order."""
    target = Path(path)
    with zipfile.ZipFile(io.BytesIO(target.read_bytes()), "r") as source:
        members = source.infolist()
        names = [member.filename for member in members]
        if len(names) != len(set(names)):
            raise ValueError("DOCX package contains duplicate member names")
        required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
        if not required.issubset(names):
            raise ValueError("DOCX package is missing required OPC members")
        entries = {member.filename: source.read(member) for member in members}

    normalized = io.BytesIO()
    with zipfile.ZipFile(
            normalized, "w", compression=zipfile.ZIP_DEFLATED,
            compresslevel=9, strict_timestamps=True) as destination:
        destination.comment = b""
        for name in sorted(entries):
            member = zipfile.ZipInfo(name, _DETERMINISTIC_DOCX_TIMESTAMP)
            member.compress_type = zipfile.ZIP_DEFLATED
            member.create_system = 3
            member.external_attr = _DETERMINISTIC_DOCX_MODE
            member.internal_attr = 0
            member.comment = b""
            member.extra = b""
            destination.writestr(
                member, entries[name], compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )

    temporary = None
    try:
        with tempfile.NamedTemporaryFile(
                mode="w+b", dir=target.parent, prefix=f".{target.name}.",
                suffix=".tmp", delete=False) as handle:
            temporary = Path(handle.name)
            handle.write(normalized.getvalue())
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary, target)
        temporary = None
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def build_keiyaku37_docx(data: dict, out_path, *, deal_type: str = "売買") -> str:
    """37条書面（契約締結時交付書面）の雛形を .docx 生成（監査KEIYAKU是正）。
    35条重説とは別の法定書面。金額・法令は確定せず、枠と確認欄を出す＝記入・記名は宅建士。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    d = data or {}
    is_sale = "賃貸" not in deal_type and "貸" not in deal_type
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Mincho"
    style.font.size = Pt(9)
    try:
        from docx.oxml.ns import qn
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Mincho")
    except Exception:
        pass
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(36)
        sec.left_margin = sec.right_margin = Pt(40)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = banner.add_run("DRAFT / 下書き・交付不可（顧客交付禁止）")
    br.bold = True
    br.font.size = Pt(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("契約書面（宅地建物取引業法第37条）　" + ("売買・交換用" if is_sale else "貸借用"))
    tr.bold = True
    tr.font.size = Pt(15)
    lead = doc.add_paragraph()
    lead.add_run("本生成物は、宅地建物取引業法第37条に基づき契約成立時に遅滞なく"
                 "交付すべき書面を作成するためのドラフトです。"
                 "この生成ドラフト自体は顧客へ交付できません。"
                 "各欄は契約内容と根拠資料に基づき確認・補正し、"
                 "宅地建物取引士が記名する必要があります。").font.size = Pt(9)

    items = [it for it in KEIYAKU37_ITEMS_COMMON
             if not (it.startswith("移転登記") and not is_sale)]
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = "Table Grid"
    for i, it in enumerate(items):
        tbl.rows[i].cells[0].text = it
        tbl.rows[i].cells[1].text = str(d.get(it) or "")

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.add_run("※ 本書は下書きの雛形です。顧客交付不可です。"
                 "金額・期日・法令の判断は確定しておらず、"
                 "契約内容と根拠資料に基づく確認・補正のうえ、宅地建物取引士が記名（押印）した"
                 "交付用書面を別途確定する必要があります。"
                 "35条重要事項説明書とは別の書面です。記載の最終責任は記名した宅地建物取引士に帰属します。"
                 ).font.size = Pt(8.5)

    out = str(Path(out_path).expanduser())
    doc.save(out)
    _normalize_docx_package(out)
    return out


def office_to_pdf(src_path, out_dir=None) -> str | None:
    """LibreOffice(soffice)で .xlsx/.docx → PDF。soffice 不在なら None。"""
    soffice = _find_soffice()
    if not soffice:
        return None
    src = Path(src_path).expanduser()
    out_dir = Path(out_dir or src.parent).expanduser()
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        command = [soffice]
        if sys.platform != "darwin":
            # macOS LibreOffice 26.x aborts when a second UserInstallation is supplied;
            # Linux/Windows retain isolation from first-run dialogs and profile locks.
            with tempfile.TemporaryDirectory(prefix="ainote_lo_profile_") as profile_dir:
                command.append(f"-env:UserInstallation={Path(profile_dir).resolve().as_uri()}")
                subprocess.run(
                    command + ["--headless", "--convert-to", "pdf", "--outdir",
                               str(out_dir), str(src)],
                    check=True, timeout=120, stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
        else:
            subprocess.run(
                command + ["--headless", "--convert-to", "pdf", "--outdir",
                           str(out_dir), str(src)],
                check=True, timeout=120, stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
    except Exception:
        return None
    pdf = out_dir / (src.stem + ".pdf")
    return str(pdf) if pdf.exists() else None


def _find_soffice():
    if str(os.environ.get("RI_HUB_DISABLE_SOFFICE") or "").strip().lower() in {
        "1", "true", "yes", "on",
    }:
        return None
    for cand in ("soffice", "libreoffice",
                 "/Applications/LibreOffice.app/Contents/MacOS/soffice"):
        if Path(cand).exists() or shutil.which(cand):
            return shutil.which(cand) or cand
    return None
